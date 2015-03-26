# -*- coding: utf-8 -*-
#登陆知乎抓取个人收藏 然后保存为word
import sys
reload(sys)  
sys.setdefaultencoding('utf-8')
import urllib
import urllib2
import cookielib
import string 
import re 
from bs4 import BeautifulSoup
from docx import Document
from docx import *
from docx.shared import Inches
from sys import exit
import socks
import socket
import os
socks.setdefaultproxy(socks.PROXY_TYPE_SOCKS5,"127.0.0.1",8088)
socket.socket =socks.socksocket

loginurl='http://www.zhihu.com/login'

headers = {'User-Agent' : 'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/34.0.1847.116 Safari/537.36',}  

postdata={
  '_xsrf':  'acab9d276ea217226d9cc94a84a231f7',
  'email':  '',
  'password': '',
  'rememberme':'y'    
}

if not os.path.exists('myimg'):
    os.mkdir('myimg')
if  os.path.exists('123.docx'):
    os.remove('123.docx')
if  os.path.exists('checkcode.gif'):
    os.remove('checkcode.gif')

mydoc=Document()
questiontitle=''
#----------------------------------------------------------------------
def  dealimg(imgcontent):
    soup=BeautifulSoup(imgcontent)
    try:
        for imglink in soup.findAll('img'):
            if imglink is not None :
                myimg= imglink.get('src')
                #print myimg
                if myimg.find('http')>=0:
                    imgsrc=urllib2.urlopen(myimg).read()
                    imgnamere=re.compile(r'http\S*/')
                    imgname=imgnamere.sub('',myimg)
                    #print imgname
                    with open(u'myimg'+'/'+imgname,'wb') as code:
                        code.write(imgsrc)
                        mydoc.add_picture(u'myimg/'+imgname,width=Inches(1.25))
    except:
        pass
    strinfo=re.compile(r'<noscript>[\s\S]*</noscript>')
    imgcontent=strinfo.sub('',imgcontent)
    strinfo=re.compile(r'<img class[\s\S]*</>')
    imgcontent=strinfo.sub('',imgcontent)
    #show all
    strinfo=re.compile(r'<a class="toggle-expand[\s\S]*</a>')
    imgcontent=strinfo.sub('',imgcontent)

    strinfo=re.compile(r'<a class=" wrap external"[\s\S]*rel="nofollow noreferrer" target="_blank">')
    imgcontent=strinfo.sub('',imgcontent)
    imgcontent=imgcontent.replace('<i class="icon-external"></i></a>','')


    imgcontent=imgcontent.replace('</b>','').replace('</p>','').replace('<p>','').replace('<p>','').replace('<br>','')
    return imgcontent
    




def enterquestionpage(pageurl):
    html=urllib2.urlopen(pageurl).read()
    soup=BeautifulSoup(html)
    questiontitle=soup.title.string
    mydoc.add_heading(questiontitle,level=3)
    for div in soup.findAll('div',{'class':'fixed-summary zm-editable-content clearfix'}):
        #print div
        conent=str(div).replace('<div class="fixed-summary zm-editable-content clearfix">','').replace('</div>','')
        
        conent=conent.decode('utf-8')
        conent=conent.replace('<br/>','\n')
        
        conent=dealimg(conent)
        ###这一块弄得太复杂了 有时间找找看有没有处理html的模块
        conent=conent.replace('<div class="fixed-summary-mask">','').replace('<blockquote>','').replace('<b>','').replace('<strong>','').replace('</strong>','').replace('<em>','').replace('</em>','').replace('</blockquote>','')
        mydoc.add_paragraph(conent,style='BodyText3')
        """file=open('222.txt','a')
        file.write(str(conent))
        file.close()"""
        

def entercollectpage(pageurl):
    html=urllib2.urlopen(pageurl).read()
    soup=BeautifulSoup(html)
    for div in soup.findAll('div',{'class':'zm-item'}):
        h2content=div.find('h2',{'class':'zm-item-title'})
        #print h2content
        if h2content is not None:
            link=h2content.find('a')
            mylink=link.get('href')
            quectionlink='http://www.zhihu.com'+mylink
            enterquestionpage(quectionlink)
            print quectionlink         



def loginzhihu():
    postdatastr=urllib.urlencode(postdata)
    '''
    cj = cookielib.LWPCookieJar()
    cookie_support = urllib2.HTTPCookieProcessor(cj)
    opener = urllib2.build_opener(cookie_support,urllib2.HTTPHandler)
    urllib2.install_opener(opener)
    '''
    h = urllib2.urlopen(loginurl) 
    request = urllib2.Request(loginurl,postdatastr,headers)
    request.get_origin_req_host
    response = urllib2.urlopen(request)
    #print response.geturl()
    text = response.read()


    collecturl='http://www.zhihu.com/collections'
    req=urllib2.urlopen(collecturl)
    if str(req.geturl())=='http://www.zhihu.com/?next=%2Fcollections':
        print 'login fail!'
        return 
    txt=req.read()

    soup=BeautifulSoup(txt)
    count=0
    divs =soup.findAll('div',{'class':'zm-item'})
    if divs is None:
        print 'login fail!'
        return
    print 'login ok!\n'
    for div in divs:
        
        link=div.find('a')
        mylink=link.get('href')
        collectlink='http://www.zhihu.com'+mylink
        entercollectpage(collectlink)
        print collectlink
        
        count+=1
        if count==1:
            return
        

def getcheckcode(thehtml):
    soup=BeautifulSoup(thehtml)
    div=soup.find('div',{'class':'js-captcha captcha-wrap'})
    if div is not None:
        #print div
        imgsrc=div.find('img')
        imglink=imgsrc.get('src')
        if imglink is not None:
            imglink='http://www.zhihu.com'+imglink

            imgcontent=urllib2.urlopen(imglink).read()
            with open('checkcode.gif','wb') as code:
                code.write(imgcontent)
            return True
        else:
            return False
    return False


if __name__=='__main__':
     
    import  getpass
    username=raw_input('input username:')
    password=getpass.getpass('Enter password: ')  
    
    postdata['email']=username
    postdata['password']=password
    postdatastr=urllib.urlencode(postdata)
    cj = cookielib.LWPCookieJar()
    cookie_support = urllib2.HTTPCookieProcessor(cj)
    opener = urllib2.build_opener(cookie_support,urllib2.HTTPHandler)
    urllib2.install_opener(opener)

    h = urllib2.urlopen(loginurl) 
    request = urllib2.Request(loginurl,postdatastr,headers)
    response = urllib2.urlopen(request)
    txt = response.read()

    if getcheckcode(txt):
        checkcode=raw_input('input checkcode:')
        postdata['captcha']=checkcode
        loginzhihu()
        mydoc.save('123.docx')
    else:
        loginzhihu()
        mydoc.save('123.docx')

    print 'the end'
    raw_input()
